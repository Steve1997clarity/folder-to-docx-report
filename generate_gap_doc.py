#!/usr/bin/env python3
"""
Generate a DOCX document explaining the gap from demo to stable production.
Target audience: non-technical supervisors and sales staff.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_normal(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def create_gap_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- HEADER LOGO ---
    header_img = os.path.join(os.getcwd(), "USER_INPUT", "DOCX_HEADER_IMAGE", "header_image.png")
    if os.path.exists(header_img):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(header_img, width=Inches(2.5))
        # Separator
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = sep.add_run("_" * 80)
        sep_run.font.color.rgb = RGBColor(200, 200, 200)
        sep_run.font.size = Pt(8)

    # --- TITLE ---
    title = doc.add_heading('From Demo to Production', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Technical Gap Analysis & Cloud Deployment Options')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()  # spacer

    # --- DOCUMENT PURPOSE ---
    add_heading_styled(doc, 'What This Document Covers', level=1)
    add_normal(doc,
        'This document explains the difference between our current live demonstration '
        'and a full production system ready for paying clients. It is written in plain '
        'language for supervisors, sales teams, and decision-makers who need to understand '
        'what is required to move from "working demo" to "reliable product".')

    # --- SECTION 1: WHAT THE DEMO DOES ---
    add_heading_styled(doc, '1. What the Demo Does Today', level=1)
    add_normal(doc,
        'The live demo allows anyone with a web link to:')
    add_bullet(doc, 'Click a button to instantly download a sample inspection report (pre-built)')
    add_bullet(doc, 'Upload their own photos and receive a professionally formatted DOCX report')
    add_bullet(doc, 'See company branding (logo, contact details) embedded in the report')
    add_normal(doc,
        'This proves the core concept works. A surveyor can turn a folder of drone photos '
        'into a formatted Word document in seconds, not hours.')

    # --- SECTION 2: DEMO LIMITATIONS ---
    add_heading_styled(doc, '2. Current Demo Limitations', level=1)
    add_normal(doc,
        'The demo runs on GitHub Codespaces \u2014 a temporary cloud workspace designed for '
        'software developers, not end users. Think of it like borrowing a friend\'s laptop '
        'to show a presentation: it works for the meeting, but you wouldn\'t run a business on it.')

    add_heading_styled(doc, 'Key Limitations:', level=2)

    # Table for limitations
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Limitation', 'What It Means', 'Impact']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    limitations = [
        ('30-Minute Auto-Shutdown',
         'The system turns off after 30 minutes of no activity',
         'Client loses access mid-demo if they step away'),
        ('No Custom Domain',
         'URL is a long, random GitHub address',
         'Looks unprofessional; hard to share with clients'),
        ('No Data Storage',
         'Uploaded files disappear when session ends',
         'Cannot save reports or track usage history'),
        ('No User Accounts',
         'Anyone with the link can access it',
         'No security, no personalisation, no usage tracking'),
        ('Server in US East Coast',
         'Data travels across the Pacific Ocean',
         '3-5 second delay for users in Hong Kong/Asia'),
        ('20 MB Upload Limit',
         'Cannot handle large inspection photo sets',
         'Real projects may have 50-200 high-res images'),
    ]

    for row_idx, (lim, meaning, impact) in enumerate(limitations, start=1):
        table.rows[row_idx].cells[0].text = lim
        table.rows[row_idx].cells[1].text = meaning
        table.rows[row_idx].cells[2].text = impact

    doc.add_paragraph()  # spacer

    # --- SECTION 3: WHAT PRODUCTION NEEDS ---
    add_heading_styled(doc, '3. What a Production System Needs', level=1)
    add_normal(doc,
        'To sell this as a reliable service to surveying firms, we need:')

    add_bullet(doc, ' \u2014 The system is always accessible, like a website (e.g., reports.metapeller.com)',
               bold_prefix='Always Online')
    add_bullet(doc, ' \u2014 Each client logs in with their own credentials',
               bold_prefix='User Accounts')
    add_bullet(doc, ' \u2014 Reports and photos are saved securely for future reference',
               bold_prefix='Data Storage')
    add_bullet(doc, ' \u2014 Server located in Hong Kong or Singapore for fast response',
               bold_prefix='Fast Response')
    add_bullet(doc, ' \u2014 Custom web address with company branding',
               bold_prefix='Professional URL')
    add_bullet(doc, ' \u2014 Handle 100+ images per report, multiple users at the same time',
               bold_prefix='Scalability')
    add_bullet(doc, ' \u2014 Encrypted data transfer and secure file storage',
               bold_prefix='Security')
    add_bullet(doc, ' \u2014 Track who uses what, when, and how often',
               bold_prefix='Usage Analytics')

    # --- SECTION 4: THE GAP ---
    add_heading_styled(doc, '4. The Gap: Demo vs Production', level=1)
    add_normal(doc,
        'Think of it like building a house. The demo is a show home \u2014 it looks right and '
        'proves the design works. But to move a family in, you need plumbing, electricity, '
        'security locks, and council approval. The "gap" is all the infrastructure that makes '
        'a working prototype into a reliable daily tool.')

    # Visual gap diagram as table
    gap_table = doc.add_table(rows=9, cols=2)
    gap_table.style = 'Light List Accent 1'
    gap_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    gap_table.rows[0].cells[0].text = 'DEMO (Show Home)'
    gap_table.rows[0].cells[1].text = 'PRODUCTION (Real Home)'
    for paragraph in gap_table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    for paragraph in gap_table.rows[0].cells[1].paragraphs:
        for run in paragraph.runs:
            run.bold = True

    gaps = [
        ('Temporary server (borrowed)', 'Dedicated server (owned/rented)'),
        ('Turns off after 30 min', 'Runs 24/7/365'),
        ('Random URL', 'Custom domain (yourcompany.com)'),
        ('No login required', 'Secure user accounts'),
        ('Files deleted on shutdown', 'Permanent cloud storage'),
        ('Single user at a time', 'Hundreds of users simultaneously'),
        ('US East server only', 'Asia-Pacific server (fast for HK)'),
        ('No monitoring', 'Performance dashboards & alerts'),
    ]

    for row_idx, (demo, prod) in enumerate(gaps, start=1):
        gap_table.rows[row_idx].cells[0].text = demo
        gap_table.rows[row_idx].cells[1].text = prod

    doc.add_paragraph()  # spacer

    # --- SECTION 5: CLOUD SOLUTIONS ---
    add_heading_styled(doc, '5. Cloud Platform Options', level=1)
    add_normal(doc,
        'To bridge this gap, we deploy the system on a major cloud platform. '
        'Below are the three leading options, compared in plain terms.')

    # --- AWS ---
    add_heading_styled(doc, 'Option A: Amazon Web Services (AWS)', level=2)
    add_normal(doc,
        'The largest cloud provider in the world. Used by Netflix, Airbnb, and most '
        'Fortune 500 companies. Has a data centre in Hong Kong.')

    p = doc.add_paragraph()
    run = p.add_run('Strengths:')
    run.bold = True
    add_bullet(doc, 'Data centre in Hong Kong \u2014 fastest possible speed for local users')
    add_bullet(doc, 'Most mature platform with widest range of services')
    add_bullet(doc, 'Strong enterprise security certifications')
    add_bullet(doc, 'Pay-as-you-go pricing \u2014 only pay for what you use')

    p = doc.add_paragraph()
    run = p.add_run('Considerations:')
    run.bold = True
    add_bullet(doc, 'More complex to set up (steeper learning curve for developers)')
    add_bullet(doc, 'Billing can be unpredictable if not monitored')
    add_bullet(doc, 'May be over-engineered for a small initial deployment')

    p = doc.add_paragraph()
    run = p.add_run('Key Services We Would Use:')
    run.bold = True
    add_bullet(doc, ' \u2014 runs the application (like a virtual computer in the cloud)', bold_prefix='EC2 or ECS')
    add_bullet(doc, ' \u2014 stores uploaded photos and generated reports permanently', bold_prefix='S3')
    add_bullet(doc, ' \u2014 stores user accounts and report records', bold_prefix='RDS')
    add_bullet(doc, ' \u2014 gives us a clean web address with security certificate', bold_prefix='CloudFront + Route 53')

    p = doc.add_paragraph()
    run = p.add_run('Estimated Monthly Cost (small scale, 10-50 users): ')
    run.bold = True
    p.add_run('HK$400 \u2013 HK$1,500/month')

    doc.add_paragraph()  # spacer

    # --- Google Cloud ---
    add_heading_styled(doc, 'Option B: Google Cloud Platform (GCP)', level=2)
    add_normal(doc,
        'Google\'s cloud offering. Used by Spotify, PayPal, and many AI/tech companies. '
        'Has data centres in Hong Kong and Taiwan.')

    p = doc.add_paragraph()
    run = p.add_run('Strengths:')
    run.bold = True
    add_bullet(doc, 'Excellent AI and machine learning integration (useful for future features)')
    add_bullet(doc, 'Simple deployment with "Cloud Run" \u2014 easiest to get started')
    add_bullet(doc, 'Generous free tier for small workloads')
    add_bullet(doc, 'Clean, developer-friendly interface')

    p = doc.add_paragraph()
    run = p.add_run('Considerations:')
    run.bold = True
    add_bullet(doc, 'Smaller market share than AWS \u2014 fewer third-party tutorials')
    add_bullet(doc, 'Enterprise support can be expensive')
    add_bullet(doc, 'Some services are newer/less battle-tested')

    p = doc.add_paragraph()
    run = p.add_run('Key Services We Would Use:')
    run.bold = True
    add_bullet(doc, ' \u2014 runs the app automatically, scales up/down with demand', bold_prefix='Cloud Run')
    add_bullet(doc, ' \u2014 stores files (like AWS S3)', bold_prefix='Cloud Storage')
    add_bullet(doc, ' \u2014 database for user accounts', bold_prefix='Cloud SQL')
    add_bullet(doc, ' \u2014 custom domain and fast content delivery', bold_prefix='Cloud CDN + DNS')

    p = doc.add_paragraph()
    run = p.add_run('Estimated Monthly Cost (small scale, 10-50 users): ')
    run.bold = True
    p.add_run('HK$300 \u2013 HK$1,200/month')

    doc.add_paragraph()  # spacer

    # --- Azure ---
    add_heading_styled(doc, 'Option C: Microsoft Azure', level=2)
    add_normal(doc,
        'Microsoft\'s cloud platform. Used by most government agencies and large corporations. '
        'Has a data centre in Hong Kong. Integrates naturally with Microsoft Office/365.')

    p = doc.add_paragraph()
    run = p.add_run('Strengths:')
    run.bold = True
    add_bullet(doc, 'Best integration with Microsoft Office \u2014 relevant for DOCX generation')
    add_bullet(doc, 'Strong presence in Hong Kong enterprise market')
    add_bullet(doc, 'Familiar to IT departments that already use Microsoft products')
    add_bullet(doc, 'Good compliance certifications for government/institutional clients')

    p = doc.add_paragraph()
    run = p.add_run('Considerations:')
    run.bold = True
    add_bullet(doc, 'Interface can be confusing for developers new to Azure')
    add_bullet(doc, 'Pricing is complex with many tiers and options')
    add_bullet(doc, 'Some services feel "enterprise-heavy" for a small startup deployment')

    p = doc.add_paragraph()
    run = p.add_run('Key Services We Would Use:')
    run.bold = True
    add_bullet(doc, ' \u2014 runs the app in a container (similar to AWS ECS)', bold_prefix='App Service')
    add_bullet(doc, ' \u2014 stores files', bold_prefix='Blob Storage')
    add_bullet(doc, ' \u2014 database', bold_prefix='Azure SQL')
    add_bullet(doc, ' \u2014 custom domain and security', bold_prefix='Front Door + DNS Zone')

    p = doc.add_paragraph()
    run = p.add_run('Estimated Monthly Cost (small scale, 10-50 users): ')
    run.bold = True
    p.add_run('HK$500 \u2013 HK$1,800/month')

    doc.add_paragraph()  # spacer

    # --- COMPARISON TABLE ---
    add_heading_styled(doc, '6. Side-by-Side Comparison', level=1)

    comp_table = doc.add_table(rows=9, cols=4)
    comp_table.style = 'Medium Shading 1 Accent 1'
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    comp_headers = ['Criteria', 'AWS', 'Google Cloud', 'Azure']
    for i, h in enumerate(comp_headers):
        comp_table.rows[0].cells[i].text = h

    comparisons = [
        ('HK Data Centre', 'Yes', 'Yes (+ Taiwan)', 'Yes'),
        ('Ease of Setup', 'Medium', 'Easy', 'Medium'),
        ('Cost (Small Scale)', 'Medium', 'Low-Medium', 'Medium-High'),
        ('AI/ML Integration', 'Good', 'Excellent', 'Good'),
        ('Microsoft Office Fit', 'Basic', 'Basic', 'Excellent'),
        ('Enterprise Credibility', 'High', 'High', 'Very High'),
        ('Scalability', 'Excellent', 'Excellent', 'Excellent'),
        ('Recommended For', 'Mature product\nwith many users', 'Fast MVP launch\nAI-heavy features', 'Corporate clients\nMS Office workflows'),
    ]

    for row_idx, (criteria, aws, gcp, azure) in enumerate(comparisons, start=1):
        comp_table.rows[row_idx].cells[0].text = criteria
        comp_table.rows[row_idx].cells[1].text = aws
        comp_table.rows[row_idx].cells[2].text = gcp
        comp_table.rows[row_idx].cells[3].text = azure

    doc.add_paragraph()  # spacer

    # --- SECTION 7: RECOMMENDATION ---
    add_heading_styled(doc, '7. Recommendation', level=1)
    add_normal(doc,
        'For the initial production launch of this document automation platform:')

    p = doc.add_paragraph()
    run = p.add_run('Recommended: Google Cloud Platform (Cloud Run)')
    run.bold = True
    run.font.size = Pt(13)

    add_normal(doc, 'Reasons:')
    add_bullet(doc, 'Fastest path from demo to production (days, not weeks)')
    add_bullet(doc, 'Lowest initial cost with generous free tier')
    add_bullet(doc, 'Built-in AI services for future intelligent features (auto-categorisation, defect detection)')
    add_bullet(doc, 'Hong Kong data centre available')
    add_bullet(doc, 'Auto-scaling: handles 1 user or 1,000 users without manual intervention')

    add_normal(doc, '')
    add_normal(doc,
        'As the platform grows and takes on enterprise/government clients, migrating to '
        'AWS or Azure for compliance reasons is straightforward \u2014 the application code '
        'remains the same, only the hosting environment changes.')

    # --- SECTION 8: TIMELINE ---
    add_heading_styled(doc, '8. Indicative Development Phases', level=1)

    phase_table = doc.add_table(rows=5, cols=3)
    phase_table.style = 'Light Shading Accent 1'
    phase_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    phase_table.rows[0].cells[0].text = 'Phase'
    phase_table.rows[0].cells[1].text = 'What Gets Built'
    phase_table.rows[0].cells[2].text = 'Outcome'
    for paragraph in phase_table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    for paragraph in phase_table.rows[0].cells[1].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    for paragraph in phase_table.rows[0].cells[2].paragraphs:
        for run in paragraph.runs:
            run.bold = True

    phases = [
        ('Phase 1:\nCloud Deployment',
         'Move app to cloud server\nCustom domain\nBasic monitoring',
         'System is "always on"\nClients can bookmark the URL'),
        ('Phase 2:\nUser Management',
         'Login system\nClient-specific branding\nReport history',
         'Each firm has their own account\nReports are saved for later'),
        ('Phase 3:\nEnterprise Features',
         'Bulk uploads (100+ images)\nTemplate library\nUsage analytics\nAPI for integration',
         'Handles real workloads\nManagement can track adoption'),
        ('Phase 4:\nAI Enhancement',
         'Auto image categorisation\nDefect detection suggestions\nSmart report formatting',
         'Further time savings\nDifferentiator vs competitors'),
    ]

    for row_idx, (phase, builds, outcome) in enumerate(phases, start=1):
        phase_table.rows[row_idx].cells[0].text = phase
        phase_table.rows[row_idx].cells[1].text = builds
        phase_table.rows[row_idx].cells[2].text = outcome

    doc.add_paragraph()  # spacer

    # --- SECTION 9: SUMMARY ---
    add_heading_styled(doc, '9. Summary for Decision Makers', level=1)

    add_normal(doc,
        'The live demo proves the concept works. Building surveyors can turn a folder of '
        'photos into a professional report with one click. The technology is validated.')

    add_normal(doc,
        'To offer this as a paid service, we need cloud infrastructure \u2014 the digital '
        'equivalent of renting proper office space instead of working from a coffee shop. '
        'This is a standard, well-understood process used by every modern software company.')

    p = doc.add_paragraph()
    run = p.add_run('Key Numbers:')
    run.bold = True

    add_bullet(doc, 'Monthly cloud hosting: HK$300 \u2013 HK$1,800 (depending on platform and scale)')
    add_bullet(doc, 'The application code already works \u2014 the investment is in infrastructure, not reinvention')
    add_bullet(doc, 'All three major cloud providers (AWS, Google, Azure) have Hong Kong data centres')
    add_bullet(doc, 'The platform can be white-labelled for each client firm')

    doc.add_paragraph()  # spacer

    # --- FOOTER CONTACT ---
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("_" * 80)
    sep_run.font.color.rgb = RGBColor(200, 200, 200)
    sep_run.font.size = Pt(8)

    footer_img = os.path.join(os.getcwd(), "USER_INPUT", "DOCX_BOTTOM_IMAGE", "bottom_image.png")
    if os.path.exists(footer_img):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(footer_img, width=Inches(4.0))

    # --- SAVE ---
    output_path = os.path.join(os.getcwd(), "Demo_to_Production_Gap_Analysis.docx")
    doc.save(output_path)
    print(f"Document saved: {output_path}")
    print(f"Size: {os.path.getsize(output_path) / 1024:.0f} KB")
    return output_path

if __name__ == '__main__':
    create_gap_document()
