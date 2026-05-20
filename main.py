import os
import json
import shutil
from io import BytesIO
from flask import Flask, request, send_file, render_template, Response
from flask_login import LoginManager, current_user
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

from feedback import db, feedback_bp, seed_users, User

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # Demo limit: 20 MB max
app.secret_key = 'secret-key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(os.getcwd(), 'data', 'feedback.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(), 'uploads')

# Init extensions
db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'feedback.login'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Register feedback blueprint
app.register_blueprint(feedback_bp)

# --- Configuration ---
MAX_IMAGES = 20          # Max images per upload (demo limit)
MAX_IMAGE_SIZE_MB = 5    # Max per-image size
THUMB_SIZE = (300, 300)  # Thumbnail dimensions

@app.errorhandler(413)
def too_large(e):
    return "File too large. Demo limit is 20 MB total.", 413

# Branding presets
BRANDING_PRESETS = {
    "hkpc": {
        "name": "HKPC",
        "description": "Hong Kong Productivity Council",
        "header": os.path.join(os.getcwd(), "BRANDING_PRESETS", "hkpc_header.jpg"),
        "footer": os.path.join(os.getcwd(), "BRANDING_PRESETS", "hkpc_footer.jpg"),
    },
    "metapeller": {
        "name": "Metapeller",
        "description": "Metapeller Limited",
        "header": os.path.join(os.getcwd(), "USER_INPUT", "DOCX_HEADER_IMAGE", "header_image.png"),
        "footer": os.path.join(os.getcwd(), "USER_INPUT", "DOCX_BOTTOM_IMAGE", "bottom_image.png"),
    },
}
DEFAULT_BRANDING = "hkpc"
DEFAULT_HEADER_IMAGE = BRANDING_PRESETS[DEFAULT_BRANDING]["header"]
DEFAULT_BOTTOM_IMAGE = BRANDING_PRESETS[DEFAULT_BRANDING]["footer"]

# Demo presets
DEMO_PRESETS = {
    "elevation_6_path_1": {
        "name": "Building Elevation Survey",
        "description": "Exterior facade inspection photos captured by drone along elevation path 1.",
        "folder": os.path.join(os.getcwd(), "Gary_Project_testset", "Elevation 6_Path 1"),
        "image_count": 8,
        "date": "August 2024",
        "output_filename": "Elevation_Survey_Report.docx",
        "preview_images": [
            "DJI_20240807104359_0322_D.JPG",
            "DJI_20240807104407_0323_D.JPG",
            "DJI_20240807104415_0324_D.JPG",
            "DJI_20240807104425_0325_D.JPG"
        ]
    },
    "path_6_rthk": {
        "name": "Roof Condition Inspection",
        "description": "Aerial roof inspection photos for ancillary building condition assessment.",
        "folder": os.path.join(os.getcwd(), "Gary_Project_testset",
                  "Path 6_RTHK_BH_Ancillary Building A_Roof_Aerial Photos"),
        "image_count": 6,
        "date": "January 2025",
        "output_filename": "Roof_Inspection_Report.docx",
        "preview_images": [
            "DJI_20250117124043_0017_V_SMALL.JPG",
            "DJI_20250117124049_0018_V_SMALL.JPG",
            "DJI_20250117124054_0019_V_SMALL.JPG",
            "DJI_20250117124058_0020_V_SMALL.JPG"
        ]
    }
}


def create_docx_with_images_header_footer(folder_path, header_image_path, bottom_image_path,
                                           output_docx="output.docx", output_folder="OUTPUT",
                                           image_width=Inches(2.2), images_per_row=3,
                                           photo_labels=None):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    sectPr = section._sectPr
    vAlign = sectPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        sectPr.append(vAlign)
    vAlign.set(qn('w:val'), 'center')

    # Header
    header_section = section.header
    header_para = header_section.paragraphs[0] if header_section.paragraphs else header_section.add_paragraph()
    header_para.text = ""
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(header_image_path):
        try:
            run = header_para.add_run()
            run.add_picture(header_image_path, width=Inches(2.5))
        except Exception as e:
            run = header_para.add_run("Error inserting header image")
            run.font.size = Pt(10)
    else:
        run = header_para.add_run("Header Image Not Found")
        run.font.size = Pt(10)

    # Body: scan for JPEG images (resize for demo to keep output small)
    valid_images = []
    for root, dirs, files in os.walk(folder_path):
        for f in sorted(files):
            if f.lower().endswith((".jpg", ".jpeg")) and not f.startswith("_custom_"):
                image_path = os.path.join(root, f)
                try:
                    with Image.open(image_path) as img:
                        # Resize to max 600px wide for demo (smaller DOCX, faster download)
                        if img.width > 450:
                            ratio = 450 / img.width
                            new_size = (450, int(img.height * ratio))
                            img = img.resize(new_size, Image.LANCZOS)
                        stream = BytesIO()
                        img.convert('RGB').save(stream, format='JPEG', quality=55)
                        stream.seek(0)
                    valid_images.append((stream, f))
                except Exception as e:
                    print(f"Skipping '{f}': {e}")
                if len(valid_images) >= MAX_IMAGES:
                    break

    table = doc.add_table(rows=0, cols=images_per_row)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(0, len(valid_images), images_per_row):
        batch = valid_images[i:i+images_per_row]
        image_row = table.add_row().cells
        for idx, (stream, img_name) in enumerate(batch):
            para = image_row[idx].paragraphs[0]
            run = para.add_run()
            try:
                run.add_picture(stream, width=image_width)
            except Exception as e:
                print(f"Error inserting '{img_name}': {e}")
        label_row = table.add_row().cells
        for idx, (stream, img_name) in enumerate(batch):
            para = label_row[idx].paragraphs[0]
            label = photo_labels.get(img_name, img_name) if photo_labels else img_name
            run = para.add_run(label)
            run.font.size = Pt(8)
            para.paragraph_format.space_after = Pt(6)

    # Footer section
    footer_section = section.footer
    for para in footer_section.paragraphs:
        p = para._element
        p.getparent().remove(p)
    # Contact image in footer (if available)
    if os.path.exists(bottom_image_path):
        footer_img_para = footer_section.add_paragraph()
        footer_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = footer_img_para.add_run()
            run.add_picture(bottom_image_path, height=Cm(1.08))
        except Exception:
            pass

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    output_path = os.path.join(output_folder, output_docx)
    doc.save(output_path)
    return output_path


# In-memory caches (generated once)
_thumb_cache = {}
_demo_report_cache = {}


def _pregenerate_demos():
    """Pre-generate demo reports on startup for instant serving."""
    print("Pre-generating demo reports...")
    output_dir = os.path.join(os.getcwd(), "OUTPUT")
    os.makedirs(output_dir, exist_ok=True)
    for preset_id, preset in DEMO_PRESETS.items():
        path = create_docx_with_images_header_footer(
            folder_path=preset["folder"],
            header_image_path=DEFAULT_HEADER_IMAGE,
            bottom_image_path=DEFAULT_BOTTOM_IMAGE,
            output_docx=preset["output_filename"],
            output_folder=output_dir
        )
        with open(path, 'rb') as f:
            _demo_report_cache[preset_id] = f.read()
        print(f"  Cached: {preset['name']} ({len(_demo_report_cache[preset_id])//1024} KB)")
    # Also pre-warm thumbnail cache
    for preset_id, preset in DEMO_PRESETS.items():
        for img_name in preset["preview_images"]:
            cache_key = f"{preset_id}/{img_name}"
            image_path = os.path.join(preset["folder"], img_name)
            if os.path.exists(image_path):
                img = Image.open(image_path)
                img.thumbnail(THUMB_SIZE)
                buf = BytesIO()
                img.convert('RGB').save(buf, format='JPEG', quality=60)
                _thumb_cache[cache_key] = buf.getvalue()
    print("Demo reports and thumbnails cached.")

@app.route('/branding/<preset_id>/<part>')
def branding_image(preset_id, part):
    """Serve branding preview images for presets."""
    if preset_id not in BRANDING_PRESETS or part not in ('header', 'footer'):
        return "Not found", 404
    image_path = BRANDING_PRESETS[preset_id][part]
    if not os.path.exists(image_path):
        return "Not found", 404
    with open(image_path, 'rb') as f:
        data = f.read()
    ext = os.path.splitext(image_path)[1].lower()
    mime = 'image/jpeg' if ext in ('.jpg', '.jpeg') else 'image/png'
    response = Response(data, mimetype=mime)
    response.headers['Cache-Control'] = 'public, max-age=86400'
    return response


@app.route('/demo/thumbnail/<preset_id>/<filename>')
def demo_thumbnail(preset_id, filename):
    if preset_id not in DEMO_PRESETS:
        return "Not found", 404
    preset = DEMO_PRESETS[preset_id]
    safe_filename = os.path.basename(filename)
    cache_key = f"{preset_id}/{safe_filename}"

    if cache_key not in _thumb_cache:
        image_path = os.path.join(preset["folder"], safe_filename)
        if not os.path.exists(image_path):
            return "Not found", 404
        img = Image.open(image_path)
        img.thumbnail(THUMB_SIZE)
        buf = BytesIO()
        img.convert('RGB').save(buf, format='JPEG', quality=60)
        _thumb_cache[cache_key] = buf.getvalue()

    response = Response(_thumb_cache[cache_key], mimetype='image/jpeg')
    response.headers['Cache-Control'] = 'public, max-age=86400'
    return response


@app.route('/demo/download/<preset_id>')
def demo_download(preset_id):
    """GET-based download (works reliably through Codespace proxy)."""
    if preset_id not in DEMO_PRESETS:
        return "Invalid preset", 404

    preset = DEMO_PRESETS[preset_id]

    # Serve from pre-generated cache (instant response)
    if preset_id in _demo_report_cache:
        data = _demo_report_cache[preset_id]
    else:
        # Fallback: generate on the fly
        output_dir = os.path.join(os.getcwd(), "OUTPUT")
        os.makedirs(output_dir, exist_ok=True)
        generated_path = create_docx_with_images_header_footer(
            folder_path=preset["folder"],
            header_image_path=DEFAULT_HEADER_IMAGE,
            bottom_image_path=DEFAULT_BOTTOM_IMAGE,
            output_docx=preset["output_filename"],
            output_folder=output_dir
        )
        with open(generated_path, 'rb') as f:
            data = f.read()

    response = Response(data, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response.headers["Content-Disposition"] = f"attachment; filename={preset['output_filename']}"
    return response


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        output_dir = os.path.join(os.getcwd(), "OUTPUT")
        if os.path.exists(output_dir):
            shutil.rmtree(output_dir)
        os.makedirs(output_dir, exist_ok=True)

        upload_dir = os.path.join(os.getcwd(), "USER_INPUT", "USER_UPLOADED_FOLDER")
        if os.path.exists(upload_dir):
            shutil.rmtree(upload_dir)
        os.makedirs(upload_dir, exist_ok=True)

        folder_files = request.files.getlist("folder_files")
        saved_count = 0
        for file in folder_files:
            if saved_count >= MAX_IMAGES:
                break
            filename = file.filename
            safe_filename = os.path.normpath(filename)
            if ".." in safe_filename:
                continue
            if not safe_filename.lower().endswith((".jpg", ".jpeg")):
                continue
            dest_path = os.path.join(upload_dir, safe_filename)
            os.makedirs(os.path.dirname(dest_path), exist_ok=True)
            file.save(dest_path)
            saved_count += 1

        # Branding selection
        branding_preset = request.form.get('branding_preset', DEFAULT_BRANDING)

        if branding_preset == 'custom':
            # Custom upload — use uploaded files, fall back to default
            header_image_path = DEFAULT_HEADER_IMAGE
            header_file = request.files.get('header_image')
            if header_file and header_file.filename:
                hname = header_file.filename.lower()
                if hname.endswith(('.png', '.jpg', '.jpeg')):
                    custom_header_path = os.path.join(upload_dir, '_custom_header' + os.path.splitext(hname)[1])
                    header_file.save(custom_header_path)
                    try:
                        with Image.open(custom_header_path) as img:
                            img.verify()
                        header_image_path = custom_header_path
                    except Exception:
                        pass

            bottom_image_path = DEFAULT_BOTTOM_IMAGE
            footer_file = request.files.get('footer_image')
            if footer_file and footer_file.filename:
                fname = footer_file.filename.lower()
                if fname.endswith(('.png', '.jpg', '.jpeg')):
                    custom_footer_path = os.path.join(upload_dir, '_custom_footer' + os.path.splitext(fname)[1])
                    footer_file.save(custom_footer_path)
                    try:
                        with Image.open(custom_footer_path) as img:
                            img.verify()
                        bottom_image_path = custom_footer_path
                    except Exception:
                        pass
        elif branding_preset in BRANDING_PRESETS:
            header_image_path = BRANDING_PRESETS[branding_preset]["header"]
            bottom_image_path = BRANDING_PRESETS[branding_preset]["footer"]
        else:
            header_image_path = DEFAULT_HEADER_IMAGE
            bottom_image_path = DEFAULT_BOTTOM_IMAGE

        # Parse photo labels from frontend
        photo_labels = None
        photo_labels_raw = request.form.get('photo_labels', '')
        if photo_labels_raw:
            try:
                photo_labels = json.loads(photo_labels_raw)
            except (json.JSONDecodeError, TypeError):
                photo_labels = None

        output_docx = "Survey_Report.docx"
        generated_docx_path = create_docx_with_images_header_footer(
            folder_path=upload_dir,
            header_image_path=header_image_path,
            bottom_image_path=bottom_image_path,
            output_docx=output_docx,
            output_folder=output_dir,
            photo_labels=photo_labels
        )

        with open(generated_docx_path, 'rb') as f:
            data = f.read()
        response = Response(
            data,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response.headers["Content-Disposition"] = f"attachment; filename={output_docx}"
        return response
    else:
        # Build demo cards
        demo_cards_html = ""
        for preset_id, preset in DEMO_PRESETS.items():
            thumbnails_html = ""
            for img_name in preset["preview_images"]:
                thumbnails_html += f'<img src="/demo/thumbnail/{preset_id}/{img_name}" alt="Survey photo" class="thumb-img" loading="lazy">'

            demo_cards_html += f'''
            <div class="col-md-6 mb-4">
                <div class="demo-card">
                    <div class="thumb-grid">
                        {thumbnails_html}
                    </div>
                    <div class="card-body">
                        <h5 class="card-title">{preset["name"]}</h5>
                        <p class="card-desc">{preset["description"]}</p>
                        <div class="card-meta">
                            <span class="badge-pill">{preset["image_count"]} photos</span>
                            <span class="badge-pill">{preset["date"]}</span>
                        </div>
                        <a href="/demo/download/{preset_id}" class="btn-demo" onclick="handleDemoClick(this)">
                            <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg>
                            Generate Sample Report
                        </a>
                    </div>
                </div>
            </div>
            '''

        return render_template('index.html', demo_cards_html=demo_cards_html)

if __name__ == '__main__':
    with app.app_context():
        os.makedirs('data', exist_ok=True)
        os.makedirs('uploads', exist_ok=True)
        db.create_all()
        seed_users()
    _pregenerate_demos()
    app.run(debug=False, port=11312, host="0.0.0.0")
